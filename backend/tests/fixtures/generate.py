"""Render the fixture personas into real files.

    uv run python tests/fixtures/generate.py

Regenerating is deterministic, so the files can be committed and diffed. Each
persona is rendered into the one format it was designed to stress — see
``Persona.stresses``.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))

from tests.fixtures.personas import (
    ALL_PERSONAS,
    FORMAT_BY_PERSONA,
    Persona,
)

OUT_DIR = Path(__file__).resolve().parent / "cvs"


# ═══════════════════════════════════════════════════════════════════════════
# Word
# ═══════════════════════════════════════════════════════════════════════════


def write_docx_single_column(persona: Persona, path: Path) -> None:
    """A classic single-column Word CV using real Word list styles."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(48)
        section.left_margin = section.right_margin = Pt(54)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(persona.name)
    run.bold = True
    run.font.size = Pt(20)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = contact.add_run(
        f"{persona.headline}  |  {persona.email}  |  {persona.phone}  |  {persona.location}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x62, 0x70)

    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = links.add_run("  ·  ".join(persona.links))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x5A, 0x62, 0x70)

    _docx_heading(doc, "SUMMARY")
    doc.add_paragraph(persona.summary)

    _docx_heading(doc, "EXPERIENCE")
    for job in persona.experience:
        head = doc.add_paragraph()
        run = head.add_run(f"{job.role}, {job.org}")
        run.bold = True
        run.font.size = Pt(11)
        # A separate run for the dates: this is exactly the run-splitting that
        # the exporter has to cope with.
        tail = head.add_run(f"    {job.location}  ·  {job.dates}")
        tail.font.size = Pt(9)
        tail.font.color.rgb = RGBColor(0x5A, 0x62, 0x70)
        for bullet in job.bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    _docx_heading(doc, "EDUCATION")
    for study in persona.education:
        para = doc.add_paragraph()
        run = para.add_run(f"{study.qualification}, {study.institution}")
        run.bold = True
        para.add_run(f"    {study.dates}").font.size = Pt(9)
        if study.detail:
            doc.add_paragraph(study.detail)

    _docx_heading(doc, "SKILLS")
    doc.add_paragraph(" · ".join(persona.skills))

    doc.save(path)


def write_docx_two_column(persona: Persona, path: Path) -> None:
    """A two-column CV laid out with a Word table — the shape that breaks
    ``document.paragraphs``, since it skips table content entirely."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt, RGBColor

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = section.bottom_margin = Pt(42)
        section.left_margin = section.right_margin = Pt(42)

    header = doc.add_paragraph()
    run = header.add_run(persona.name)
    run.bold = True
    run.font.size = Pt(22)
    sub = doc.add_paragraph()
    run = sub.add_run(persona.headline)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x14, 0x65, 0x5C)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    sidebar, main = table.rows[0].cells
    sidebar.width = Pt(150)
    main.width = Pt(340)

    _cell_heading(sidebar, "CONTACT")
    for line in (persona.email, persona.phone, persona.location, *persona.links):
        sidebar.add_paragraph(line).runs[0].font.size = Pt(9)

    _cell_heading(sidebar, "SKILLS")
    for skill in persona.skills:
        sidebar.add_paragraph(skill, style="List Bullet").runs[0].font.size = Pt(9)

    _cell_heading(sidebar, "EDUCATION")
    for study in persona.education:
        para = sidebar.add_paragraph()
        run = para.add_run(study.qualification)
        run.bold = True
        run.font.size = Pt(9)
        detail = sidebar.add_paragraph(f"{study.institution}, {study.dates}")
        detail.runs[0].font.size = Pt(9)

    _cell_heading(main, "PROFILE", first=True)
    main.add_paragraph(persona.summary)

    _cell_heading(main, "EXPERIENCE")
    for job in persona.experience:
        head = main.add_paragraph()
        run = head.add_run(f"{job.role} — {job.org}")
        run.bold = True
        run.font.size = Pt(11)
        meta = main.add_paragraph(f"{job.location}  ·  {job.dates}")
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x5A, 0x62, 0x70)
        for bullet in job.bullets:
            main.add_paragraph(bullet, style="List Bullet")

    doc.save(path)


def _docx_heading(doc: object, text: str) -> None:
    from docx.shared import Pt, RGBColor

    para = doc.add_paragraph()  # type: ignore[attr-defined]
    para.paragraph_format.space_before = Pt(11)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x14, 0x65, 0x5C)


def _cell_heading(cell: object, text: str, *, first: bool = False) -> None:
    from docx.shared import Pt, RGBColor

    # The first paragraph of a fresh cell already exists and is empty; reuse it
    # so the document does not start with a stray blank line.
    para = cell.paragraphs[0] if first and not cell.paragraphs[0].text else cell.add_paragraph()  # type: ignore[attr-defined]
    para.paragraph_format.space_before = Pt(9)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x14, 0x65, 0x5C)


# ═══════════════════════════════════════════════════════════════════════════
# PDF
# ═══════════════════════════════════════════════════════════════════════════

_INK = (0.086, 0.094, 0.114)
_SLATE = (0.353, 0.384, 0.439)
_TEAL = (0.078, 0.396, 0.361)


def write_pdf_single_column(persona: Persona, path: Path) -> None:
    """A single-column PDF with mixed fonts, sizes and an accent colour."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    width, height = A4
    c = canvas.Canvas(str(path), pagesize=A4)
    left, right = 56, width - 56
    y = height - 62

    c.setFillColorRGB(*_INK)
    c.setFont("Helvetica-Bold", 21)
    c.drawString(left, y, persona.name)
    y -= 17

    c.setFillColorRGB(*_TEAL)
    c.setFont("Helvetica", 11.5)
    c.drawString(left, y, persona.headline)
    y -= 15

    c.setFillColorRGB(*_SLATE)
    c.setFont("Helvetica", 8.5)
    c.drawString(left, y, f"{persona.email}  ·  {persona.phone}  ·  {persona.location}")
    y -= 11
    c.drawString(left, y, "  ·  ".join(persona.links))
    y -= 22

    y = _pdf_heading(c, "SUMMARY", left, y)
    y = _pdf_wrap(c, persona.summary, left, y, right - left, "Helvetica", 10)
    y -= 10

    y = _pdf_heading(c, "EXPERIENCE", left, y)
    for job in persona.experience:
        c.setFillColorRGB(*_INK)
        c.setFont("Helvetica-Bold", 10.5)
        c.drawString(left, y, f"{job.role}, {job.org}")
        c.setFillColorRGB(*_SLATE)
        c.setFont("Helvetica-Oblique", 8.5)
        c.drawRightString(right, y, f"{job.location}  ·  {job.dates}")
        y -= 13
        for bullet in job.bullets:
            y = _pdf_bullet(c, bullet, left, y, right - left)
        y -= 7
        if y < 90:
            c.showPage()
            y = height - 62

    y = _pdf_heading(c, "EDUCATION", left, y)
    for study in persona.education:
        c.setFillColorRGB(*_INK)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(left, y, f"{study.qualification}, {study.institution}")
        c.setFillColorRGB(*_SLATE)
        c.setFont("Helvetica", 8.5)
        c.drawRightString(right, y, study.dates)
        y -= 14

    y -= 6
    y = _pdf_heading(c, "SKILLS", left, y)
    _pdf_wrap(c, " · ".join(persona.skills), left, y, right - left, "Helvetica", 10)

    c.save()


def write_pdf_two_column(persona: Persona, path: Path) -> None:
    """A sidebar-plus-main PDF. Reading order is the whole test here: a naive
    top-to-bottom sweep interleaves the two columns into nonsense."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    width, height = A4
    c = canvas.Canvas(str(path), pagesize=A4)

    gutter_x = 190.0
    side_l, side_r = 46.0, gutter_x - 24
    main_l, main_r = gutter_x + 10, width - 46

    c.setFillColorRGB(0.957, 0.961, 0.957)
    c.rect(0, 0, gutter_x - 10, height, fill=1, stroke=0)

    # ── sidebar ──────────────────────────────────────────────────────────
    y = height - 64
    c.setFillColorRGB(*_INK)
    c.setFont("Helvetica-Bold", 17)
    for word in persona.name.split():
        c.drawString(side_l, y, word)
        y -= 19
    y -= 2
    c.setFillColorRGB(*_TEAL)
    c.setFont("Helvetica", 9.5)
    c.drawString(side_l, y, persona.headline)
    y -= 26

    y = _pdf_heading(c, "CONTACT", side_l, y, size=9)
    c.setFillColorRGB(*_SLATE)
    c.setFont("Helvetica", 8)
    for line in (persona.email, persona.phone, persona.location, *persona.links):
        y = _pdf_wrap(c, line, side_l, y, side_r - side_l, "Helvetica", 8, leading=11)
    y -= 12

    y = _pdf_heading(c, "SKILLS", side_l, y, size=9)
    c.setFillColorRGB(*_SLATE)
    c.setFont("Helvetica", 8)
    for skill in persona.skills:
        c.drawString(side_l, y, f"· {skill}")
        y -= 11
    y -= 12

    y = _pdf_heading(c, "EDUCATION", side_l, y, size=9)
    for study in persona.education:
        c.setFillColorRGB(*_INK)
        c.setFont("Helvetica-Bold", 8.5)
        y = _pdf_wrap(
            c, study.qualification, side_l, y, side_r - side_l, "Helvetica-Bold", 8.5, leading=11
        )
        c.setFillColorRGB(*_SLATE)
        c.setFont("Helvetica", 8)
        y = _pdf_wrap(
            c,
            f"{study.institution}, {study.dates}",
            side_l,
            y,
            side_r - side_l,
            "Helvetica",
            8,
            leading=11,
        )
        y -= 8

    # ── main column ──────────────────────────────────────────────────────
    y = height - 64
    y = _pdf_heading(c, "PROFILE", main_l, y)
    y = _pdf_wrap(c, persona.summary, main_l, y, main_r - main_l, "Helvetica", 9.5)
    y -= 12

    y = _pdf_heading(c, "EXPERIENCE", main_l, y)
    for job in persona.experience:
        c.setFillColorRGB(*_INK)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(main_l, y, f"{job.role} — {job.org}")
        y -= 12
        c.setFillColorRGB(*_SLATE)
        c.setFont("Helvetica-Oblique", 8.5)
        c.drawString(main_l, y, f"{job.location}  ·  {job.dates}")
        y -= 13
        for bullet in job.bullets:
            y = _pdf_bullet(c, bullet, main_l, y, main_r - main_l, size=9.5)
        y -= 8

    c.save()


def _pdf_heading(c: object, text: str, x: float, y: float, *, size: float = 10.5) -> float:
    c.setFillColorRGB(*_TEAL)  # type: ignore[attr-defined]
    c.setFont("Helvetica-Bold", size)  # type: ignore[attr-defined]
    c.drawString(x, y, text)  # type: ignore[attr-defined]
    return y - (size + 6)


def _pdf_bullet(
    c: object, text: str, x: float, y: float, width: float, *, size: float = 10
) -> float:
    c.setFillColorRGB(*_INK)  # type: ignore[attr-defined]
    c.setFont("Helvetica", size)  # type: ignore[attr-defined]
    c.drawString(x, y, "•")  # type: ignore[attr-defined]
    return _pdf_wrap(c, text, x + 12, y, width - 12, "Helvetica", size)


def _pdf_wrap(
    c: object,
    text: str,
    x: float,
    y: float,
    width: float,
    font: str,
    size: float,
    *,
    leading: float | None = None,
) -> float:
    from reportlab.pdfbase.pdfmetrics import stringWidth

    step = leading if leading is not None else size + 3.4
    c.setFont(font, size)  # type: ignore[attr-defined]
    words, line = text.split(), ""
    for word in words:
        candidate = f"{line} {word}".strip()
        if stringWidth(candidate, font, size) <= width:
            line = candidate
            continue
        c.drawString(x, y, line)  # type: ignore[attr-defined]
        y -= step
        line = word
    if line:
        c.drawString(x, y, line)  # type: ignore[attr-defined]
        y -= step
    return y


# ═══════════════════════════════════════════════════════════════════════════
# LaTeX and plain text
# ═══════════════════════════════════════════════════════════════════════════


def _tex_escape(text: str) -> str:
    out = text
    for char in ("&", "%", "$", "#", "_"):
        out = out.replace(char, "\\" + char)
    return out


def write_tex(persona: Persona, path: Path) -> None:
    lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[margin=0.75in]{geometry}",
        r"\usepackage[scaled]{helvet}",
        r"\renewcommand{\familydefault}{\sfdefault}",
        r"\usepackage{enumitem}",
        r"\usepackage[hidelinks]{hyperref}",
        r"\usepackage{titlesec}",
        r"\titleformat{\section}{\large\bfseries}{}{0em}{}[\titlerule]",
        r"\setlist[itemize]{leftmargin=*,itemsep=1pt,topsep=2pt}",
        r"\pagestyle{empty}",
        "",
        r"\begin{document}",
        "",
        r"\begin{center}",
        rf"{{\LARGE \textbf{{{_tex_escape(persona.name)}}}}} \\[3pt]",
        rf"{_tex_escape(persona.headline)} \\[2pt]",
        rf"{_tex_escape(persona.email)} $\cdot$ {_tex_escape(persona.phone)} "
        rf"$\cdot$ {_tex_escape(persona.location)} \\",
        rf"{' $\\cdot$ '.join(_tex_escape(link) for link in persona.links)}",
        r"\end{center}",
        "",
        r"\section{Summary}",
        _tex_escape(persona.summary),
        "",
        r"\section{Experience}",
    ]

    for job in persona.experience:
        lines += [
            rf"\textbf{{{_tex_escape(job.role)}}}, {_tex_escape(job.org)} "
            rf"\hfill {_tex_escape(job.dates)}",
            "",
            r"\begin{itemize}",
        ]
        lines += [rf"  \item {_tex_escape(bullet)}" for bullet in job.bullets]
        lines += [r"\end{itemize}", ""]

    lines += [r"\section{Education}"]
    for study in persona.education:
        lines += [
            rf"\textbf{{{_tex_escape(study.qualification)}}}, "
            rf"{_tex_escape(study.institution)} \hfill {_tex_escape(study.dates)}",
            "",
        ]

    if persona.projects:
        lines += [r"\section{Projects}", r"\begin{itemize}"]
        lines += [
            rf"  \item \textbf{{{_tex_escape(title)}}} --- {_tex_escape(detail)}"
            for title, detail in persona.projects
        ]
        lines += [r"\end{itemize}", ""]

    lines += [
        r"\section{Skills}",
        _tex_escape(" $\\cdot$ ".join(persona.skills)).replace(r"\$\cdot\$", r"$\cdot$"),
        "",
        r"\end{document}",
        "",
    ]
    path.write_text("\n".join(lines), encoding="utf-8")


def write_txt(persona: Persona, path: Path) -> None:
    lines = [
        persona.name,
        persona.headline,
        f"{persona.email} | {persona.phone} | {persona.location}",
        " | ".join(persona.links),
        "",
        "SUMMARY",
        persona.summary,
        "",
        "EXPERIENCE",
        "",
    ]
    for job in persona.experience:
        lines += [f"{job.role}, {job.org} ({job.location}) — {job.dates}"]
        lines += [f"- {bullet}" for bullet in job.bullets]
        lines += [""]

    lines += ["EDUCATION", ""]
    for study in persona.education:
        lines += [f"{study.qualification}, {study.institution} — {study.dates}"]
        if study.detail:
            lines += [f"  {study.detail}"]
    lines += ["", "SKILLS", ", ".join(persona.skills), ""]
    path.write_text("\n".join(lines), encoding="utf-8")


# ═══════════════════════════════════════════════════════════════════════════
# Entry point
# ═══════════════════════════════════════════════════════════════════════════

_TWO_COLUMN = {"daniel_okonkwo_swe", "sofia_ramos_design"}


def generate_all(out_dir: Path = OUT_DIR) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    written: list[Path] = []

    for persona in ALL_PERSONAS:
        fmt = FORMAT_BY_PERSONA[persona.key]
        path = out_dir / f"{persona.key}.{fmt}"
        two_column = persona.key in _TWO_COLUMN

        if fmt == "docx":
            (write_docx_two_column if two_column else write_docx_single_column)(persona, path)
        elif fmt == "pdf":
            (write_pdf_two_column if two_column else write_pdf_single_column)(persona, path)
        elif fmt == "tex":
            write_tex(persona, path)
        else:
            write_txt(persona, path)

        written.append(path)
        print(f"  {path.name:34} {persona.stresses}")

    return written


if __name__ == "__main__":
    print("Generating fixture CVs into", OUT_DIR)
    generate_all()
    print("done.")
