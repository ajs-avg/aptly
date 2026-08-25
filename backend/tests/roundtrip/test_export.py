"""Export tests — the promise that you get your own file back.

Two claims are under test here, and they are the whole reason the source-anchor
machinery exists:

1. **Untouched in, untouched out.** Exporting a CV nobody edited returns the
   original bytes. No re-encoding, no reformatting, no silent drift.
2. **An edit changes the edited line and nothing else.** For .docx that means
   every other run keeps its own bold, size and colour, because it was never
   written to.

PDF is held to a different, honestly weaker standard: it cannot be edited in
place, so it is rebuilt, and the test asserts that no content is lost and that
the result is flagged as a rebuild.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from aptly.export import changed_nodes, export_cv
from aptly.ingest import parse_cv, parse_pasted
from aptly.model.anchors import SyntheticAnchor, is_writable
from aptly.model.document import CVDocument, TextNode

from tests.fixtures.personas import ALL_PERSONAS, FORMAT_BY_PERSONA, Persona

CV_DIR = Path(__file__).resolve().parent.parent / "fixtures" / "cvs"

#: A CV that arrived as text in a textarea rather than as a file. Written out
#: here rather than loaded from `fixtures/cvs`, because the whole point of these
#: cases is that there is no file.
PASTED_CV = """Aman Mishra
aman@example.com | +91 98765 43210 | Bengaluru

SUMMARY
Product manager with six years across hardware and software launches.

EXPERIENCE
Senior Product Manager, Kalyra - 2021 to present
- Cut new-site ramp time from 12 weeks to 6 by rebuilding the onboarding flow.
- Ran discovery with 40 customers and shipped a pricing change worth 8% ARR.

SKILLS
Python, SQL, RAG, roadmapping, discovery

EDUCATION
B.Tech, Computer Science, VIT - 2018
"""

#: Formats we edit in place, and therefore hold to byte-identical round-tripping.
EDITABLE_FORMATS = {"docx", "tex", "txt", "md"}

NEW_TEXT = "Rewrote this line to match the job post, using only what was already true."


def _path(persona: Persona) -> Path:
    path = CV_DIR / f"{persona.key}.{FORMAT_BY_PERSONA[persona.key]}"
    if not path.exists():
        pytest.skip("fixture missing: run `uv run python tests/fixtures/generate.py`")
    return path


def _load(persona: Persona) -> tuple[bytes, CVDocument]:
    path = _path(persona)
    data = path.read_bytes()
    return data, parse_cv(data, path.name)


def _ids(personas: tuple[Persona, ...]) -> list[str]:
    return [p.key for p in personas]


_IN_PLACE = tuple(p for p in ALL_PERSONAS if FORMAT_BY_PERSONA[p.key] in EDITABLE_FORMATS)
_REBUILT = tuple(p for p in ALL_PERSONAS if FORMAT_BY_PERSONA[p.key] == "pdf")


# ═══════════════════════════════════════════════════════════════════════════
# Nothing changed means nothing changes
# ═══════════════════════════════════════════════════════════════════════════


@pytest.mark.parametrize("persona", _IN_PLACE, ids=_ids(_IN_PLACE))
def test_untouched_export_is_byte_identical(persona: Persona) -> None:
    original, document = _load(persona)
    assert export_cv(original, document).data == original


@pytest.mark.parametrize("persona", ALL_PERSONAS, ids=_ids(ALL_PERSONAS))
def test_untouched_document_reports_no_changed_nodes(persona: Persona) -> None:
    original, document = _load(persona)
    assert changed_nodes(original, document) == []


@pytest.mark.parametrize("persona", ALL_PERSONAS, ids=_ids(ALL_PERSONAS))
def test_one_edit_is_reported_as_one_change(persona: Persona) -> None:
    original, document = _load(persona)
    target = document.editable_nodes[0]
    document.apply(target.id, target.text, NEW_TEXT)

    changed = changed_nodes(original, document)
    assert [node.id for node in changed] == [target.id]


# ═══════════════════════════════════════════════════════════════════════════
# In-place editing
# ═══════════════════════════════════════════════════════════════════════════


@pytest.mark.parametrize("persona", _IN_PLACE, ids=_ids(_IN_PLACE))
def test_edit_lands_in_the_exported_file(persona: Persona) -> None:
    original, document = _load(persona)
    target = next(node for node in document.editable_nodes if node.role == "bullet")
    document.apply(target.id, target.text, NEW_TEXT)

    reparsed = parse_cv(export_cv(original, document).data, _path(persona).name)
    assert NEW_TEXT in reparsed.plain_text()


@pytest.mark.parametrize("persona", _IN_PLACE, ids=_ids(_IN_PLACE))
def test_edit_leaves_every_other_line_alone(persona: Persona) -> None:
    original, document = _load(persona)
    target = next(node for node in document.editable_nodes if node.role == "bullet")
    before = {node.id: node.text for node in document.nodes if node.id != target.id}

    document.apply(target.id, target.text, NEW_TEXT)
    reparsed = parse_cv(export_cv(original, document).data, _path(persona).name)

    after = {node.id: node.text for node in reparsed.nodes if node.id != target.id}
    shared = before.keys() & after.keys()
    assert shared, "expected node ids to survive the round trip"
    assert {k: before[k] for k in shared} == {k: after[k] for k in shared}


def test_docx_edit_preserves_run_formatting() -> None:
    """The core .docx claim: rewriting a bullet must not flatten the document's
    fonts, sizes or colours. We compare every run's formatting before and after,
    excluding only the paragraph that was deliberately changed."""
    from docx import Document as open_docx

    from tests.fixtures.personas import PRIYA

    original, document = _load(PRIYA)
    target = next(node for node in document.editable_nodes if node.role == "bullet")
    edited_paragraph = target.anchor.paragraph_index  # type: ignore[union-attr]

    document.apply(target.id, target.text, NEW_TEXT)
    exported = export_cv(original, document).data

    before = _run_formatting(original, skip=edited_paragraph)
    after = _run_formatting(exported, skip=edited_paragraph)
    assert before == after

    # And the intended change really is present.
    assert NEW_TEXT in "\n".join(p.text for p in open_docx(BytesIO(exported)).paragraphs)


def test_docx_edit_touches_only_one_run() -> None:
    """Changing a single word should rewrite one run, not rebuild the paragraph."""
    from docx import Document as open_docx
    from docx.shared import Pt

    buffer = BytesIO()
    doc = open_docx()
    paragraph = doc.add_paragraph()
    first = paragraph.add_run("Cut ramp time ")
    first.bold = True
    paragraph.add_run("by half")
    tail = paragraph.add_run(" across four markets.")
    tail.font.size = Pt(9)
    doc.save(buffer)

    from aptly.export.runs import replace_span

    reopened = open_docx(BytesIO(buffer.getvalue()))
    target = reopened.paragraphs[0]
    full = target.text
    start = full.index("by half")
    assert replace_span(target, start, start + len("by half"), "by 50%")

    assert target.text == "Cut ramp time by 50% across four markets."
    assert target.runs[0].text == "Cut ramp time "
    assert target.runs[0].bold is True
    assert target.runs[2].text == " across four markets."
    assert target.runs[2].font.size == Pt(9)


def test_tex_edit_is_escaped_back_into_source() -> None:
    """Text destined for LaTeX must be re-escaped, or a stray % comments out
    the rest of the line and silently deletes content."""
    from tests.fixtures.personas import ARJUN

    original, document = _load(ARJUN)
    target = next(node for node in document.editable_nodes if node.role == "bullet")
    document.apply(target.id, target.text, "Improved recall by 30% & cut cost_per_query.")

    source = export_cv(original, document).data.decode("utf-8")
    assert r"30\%" in source
    assert r"\&" in source
    assert r"cost\_per\_query" in source

    reparsed = parse_cv(source.encode("utf-8"), "arjun_patel_ml.tex")
    assert "Improved recall by 30% & cut cost_per_query." in reparsed.plain_text()


# ═══════════════════════════════════════════════════════════════════════════
# PDF: rebuilt, and honest about it
# ═══════════════════════════════════════════════════════════════════════════


@pytest.mark.parametrize("persona", _REBUILT, ids=_ids(_REBUILT))
def test_pdf_export_is_flagged_as_a_rebuild(persona: Persona) -> None:
    original, document = _load(persona)
    result = export_cv(original, document)

    assert result.rebuilt is True
    assert result.notes, "a rebuild must be explained to the user, never silent"
    assert result.data.startswith(b"%PDF")


@pytest.mark.parametrize("persona", _REBUILT, ids=_ids(_REBUILT))
def test_pdf_rebuild_loses_no_achievements(persona: Persona) -> None:
    """A rebuild may move pixels. It may not drop sentences."""
    original, document = _load(persona)
    rebuilt = export_cv(original, document).data
    reparsed = parse_cv(rebuilt, f"{persona.key}.pdf")

    text = _squash(reparsed.plain_text())
    for job in persona.experience:
        for bullet in job.bullets:
            assert _squash(bullet) in text, f"rebuild lost: {bullet[:60]}…"


@pytest.mark.parametrize("persona", _REBUILT, ids=_ids(_REBUILT))
def test_pdf_rebuild_keeps_the_structure(persona: Persona) -> None:
    original, document = _load(persona)
    reparsed = parse_cv(export_cv(original, document).data, f"{persona.key}.pdf")

    assert reparsed.contact.name == persona.name
    experience = reparsed.section("experience")
    assert experience is not None
    assert len(experience.entries) == len(persona.experience)


@pytest.mark.parametrize("persona", _REBUILT, ids=_ids(_REBUILT))
def test_pdf_warns_on_ingest_that_it_cannot_be_edited(persona: Persona) -> None:
    """The user learns about the rebuild when they upload, not when they
    download and find their layout changed."""
    _, document = _load(persona)
    assert any("cannot be edited" in warning for warning in document.warnings)
    assert document.style_profile.inferred is True


def test_typst_special_characters_survive_a_rebuild() -> None:
    """A CV mentioning C#, $ or _ must not turn into Typst markup."""
    from aptly.ingest import parse_pasted

    document = parse_pasted(
        "Ada Lovelace\nada@example.com\n\nSKILLS\nC#, F#, $LATEX, snake_case, a*b, <html>\n"
    )
    document.source_format = "pdf"  # type: ignore[assignment]
    result = export_cv(b"", document)

    assert result.data.startswith(b"%PDF")


# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════


def _run_formatting(data: bytes, *, skip: int) -> list[tuple]:
    """Every run's visual properties, keyed by position, minus one paragraph."""
    from aptly.ingest.docx import walk_paragraphs
    from docx import Document as open_docx

    handle = open_docx(BytesIO(data))
    out: list[tuple] = []
    for index, paragraph, table_path in walk_paragraphs(handle):
        if index == skip:
            continue
        for run_index, run in enumerate(paragraph.runs):
            out.append(
                (
                    index,
                    table_path,
                    run_index,
                    run.text,
                    run.bold,
                    run.italic,
                    run.underline,
                    run.font.size,
                    run.font.name,
                    run.font.color.rgb if run.font.color and run.font.color.rgb else None,
                )
            )
    return out


def _squash(text: str) -> str:
    return "".join(ch for ch in text.lower() if ch.isalnum())


# ═══════════════════════════════════════════════════════════════════════════
# Nothing to edit in place
#
# Two ordinary flows post an edited document with no original bytes behind it,
# and both used to be handled as though there were:
#
#   - a **pasted** CV, which never had a file at all;
#   - the **rebuilt** CV, which the browser deliberately does not send the
#     uploaded file for, because the second document is not that file.
#
# The edit path parsed the empty bytes, concluded nothing had changed, and
# returned them — so the download was a zero-byte file for .txt, .md and .tex,
# and a `BadZipFile` crash for .docx. Silence and a crash, for the same cause.
# ═══════════════════════════════════════════════════════════════════════════


@pytest.mark.parametrize("fmt", sorted(EDITABLE_FORMATS | {"pdf"}))
def test_export_without_the_original_produces_a_real_file(fmt: str) -> None:
    document = parse_pasted(PASTED_CV)
    document.source_format = fmt  # type: ignore[assignment]
    document.source_filename = f"cv.{fmt}"

    result = export_cv(b"", document, fmt)

    assert result.data, f"{fmt} exported an empty file"
    assert len(result.data) > 200


@pytest.mark.parametrize("target", ["docx", "pdf", "tex", "md", "txt", None])
def test_a_pasted_cv_downloads_in_every_format(target: str | None) -> None:
    document = parse_pasted(PASTED_CV)

    result = export_cv(b"", document, target)

    assert result.data
    # There was never a file to preserve, so this is a new document and the
    # response has to say so — the UI shows that note to the person.
    assert result.rebuilt
    assert result.notes


def test_export_without_the_original_keeps_the_content() -> None:
    document = parse_pasted(PASTED_CV)

    text = export_cv(b"", document, "txt").data.decode("utf-8")

    assert "Kalyra" in text
    assert "12 weeks to 6" in text


# ═══════════════════════════════════════════════════════════════════════════
# A claimed line survives the trip back
#
# The browser owns the document for the whole editing session and posts it back
# at export, so it is the other author of this model. When somebody adds a line
# through the skill-gap flow it arrives carrying a synthetic anchor whose origin
# the server had never been taught — and the export refused the whole document
# with "Aptly could not read the edited CV", at the moment they were trying to
# download their work.
# ═══════════════════════════════════════════════════════════════════════════


def test_a_claimed_line_validates_and_exports() -> None:
    document = parse_pasted(PASTED_CV)
    section = next(s for s in document.sections if s.kind == "skills")
    section.loose_nodes.append(
        TextNode(
            id="claim_1a2b3c",
            role="skill_line",
            text="Ran the nightly deploy on Kubernetes at Kalyra, across three environments.",
            anchor=SyntheticAnchor(origin="claim"),
        )
    )

    # Exactly what the browser posts, and exactly how the endpoint reads it.
    revalidated = CVDocument.model_validate_json(document.model_dump_json())

    assert "Kubernetes" in export_cv(b"", revalidated, "txt").data.decode("utf-8")


def test_a_claimed_line_is_never_written_through_to_the_original() -> None:
    # It has no address in the uploaded file, because it was never in it.
    assert not is_writable(SyntheticAnchor(origin="claim"))
