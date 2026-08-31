"""Writing a CV out in a format it did not arrive in.

The rest of :mod:`aptly.export` exists to *edit* — to put changed lines back into
the person's own file so their formatting survives untouched. That is the right
default and it is the product's strongest guarantee, but it only answers the
question "give me my file back". It cannot answer "I uploaded a .docx and the
application form wants a PDF", which people hit constantly and which currently
means opening Word.

So this module does the other thing: it builds a *fresh* document in whichever
format was asked for, from the canonical model.

The distinction matters enough to be visible in the UI. Downloading the source
format is lossless — byte-identical when nothing changed. Downloading any other
format is a rebuild: the content is exactly right, the layout is Aptly's rather
than the original's, and the person is told so before they click.

PDF already had a renderer, because a PDF can never be edited in place. This
adds the other three, so every format can be reached from every format.
"""

from __future__ import annotations

import io

from aptly.model.document import CVDocument

#: What a rebuild costs, said plainly. Shown wherever a target format differs
#: from the source, because "your formatting is gone" is not something to
#: discover after downloading.
REBUILD_NOTE = (
    "This is a new document rather than an edit of your file, so it uses Aptly's "
    "layout rather than your original's. The wording is exactly what you approved."
)


def render(document: CVDocument, target: str, tex_renderer: str = "stock") -> bytes:
    """Build ``document`` afresh in ``target`` format.

    ``tex_renderer`` names a LaTeX writer for templates whose layout lives in
    macros rather than in a style profile — a title and a date on one ruled
    row, small caps under a rule — none of which is a font size and so none of
    which a :class:`StyleProfile` can carry.
    """
    if target == "pdf":
        from aptly.export.pdf import rebuild_pdf

        return rebuild_pdf(document, filename="cv.pdf").data
    if target == "docx":
        return render_docx(document)
    if target == "tex":
        if tex_renderer == "format1":
            from aptly.export.latex_format1 import render_format1

            return render_format1(document).encode("utf-8")
        return render_tex(document).encode("utf-8")
    if target == "md":
        return render_markdown(document).encode("utf-8")
    if target == "txt":
        return render_text(document).encode("utf-8")

    from aptly.errors import UnsupportedFormatError

    raise UnsupportedFormatError(
        f"Aptly cannot write {target} files.",
        hint="Choose .docx, .pdf, .tex, .md or plain text.",
    )


# ═══════════════════════════════════════════════════════════════════════════
# Word
# ═══════════════════════════════════════════════════════════════════════════


def render_docx(document: CVDocument) -> bytes:
    """A fresh .docx, styled from the profile the parser measured.

    Built rather than templated: carrying a .docx template in the repo would mean
    every rebuild looked like that template regardless of what the person
    uploaded, and the point of the style profile is that a rebuild can look like
    *their* CV.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    style = document.style_profile
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(style.margins.top_pt)
    section.bottom_margin = Pt(style.margins.bottom_pt)
    section.left_margin = Pt(style.margins.left_pt)
    section.right_margin = Pt(style.margins.right_pt)

    normal = doc.styles["Normal"]
    normal.font.name = style.body.family
    normal.font.size = Pt(style.body.size_pt)

    def write(
        text: str, spec, *, centre: bool = False, bullet: bool = False, space_before: float = 0.0
    ):
        if not text.strip():
            return
        paragraph = doc.add_paragraph(style="List Bullet" if bullet else None)
        if centre:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(space_before)
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text)
        run.font.name = spec.family
        run.font.size = Pt(spec.size_pt)
        run.bold = spec.bold
        run.italic = spec.italic
        if spec.color and spec.color != "#000000":
            run.font.color.rgb = RGBColor.from_string(spec.color.lstrip("#").upper())

    contact = document.contact
    if contact.name:
        write(contact.name, style.name, centre=True)
    details = [d for d in (contact.email, contact.phone, contact.location) if d]
    if details:
        write("  ·  ".join(details), style.body, centre=True)
    if contact.links:
        write("  ·  ".join(contact.links), style.body, centre=True)

    for block in document.sections:
        if block.kind == "header":
            continue
        if block.title:
            write(
                _heading_text(block.title, style.heading_transform),
                style.section_heading,
                space_before=style.heading_space_before_pt,
            )
        for node in block.loose_nodes:
            write(node.text, style.body)
        for entry in block.entries:
            heading = _entry_heading(entry)
            if heading:
                write(heading, style.entry_heading, space_before=6)
            for node in entry.bullets:
                write(node.text, style.body, bullet=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════
# LaTeX
# ═══════════════════════════════════════════════════════════════════════════

#: The characters LaTeX treats as syntax. Left unescaped, a CV containing "R&D"
#: or "100%" produces a document that will not compile — and the person only
#: finds out when their build fails, long after they have left.
_TEX_ESCAPES = {
    "\\": r"\textbackslash{}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "{": r"\{",
    "}": r"\}",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def _tex(text: str) -> str:
    return "".join(_TEX_ESCAPES.get(character, character) for character in text)


def render_tex(document: CVDocument) -> str:
    """A self-contained LaTeX source that compiles with a stock article class.

    No custom class and no external package beyond what TeX Live ships by
    default: a .tex file that needs something installed to build is not a
    deliverable, it is homework.
    """
    style = document.style_profile
    contact = document.contact
    margin = round(style.margins.left_pt / 72, 2)

    out = [
        "\\documentclass[11pt,a4paper]{article}",
        "\\usepackage[T1]{fontenc}",
        "\\usepackage[utf8]{inputenc}",
        f"\\usepackage[margin={margin}in]{{geometry}}",
        "\\usepackage{enumitem}",
        "\\usepackage{titlesec}",
        "\\usepackage[hidelinks]{hyperref}",
        "\\pagestyle{empty}",
        "\\setlist[itemize]{leftmargin=*,itemsep=1pt,topsep=2pt}",
        "\\titleformat{\\section}{\\large\\bfseries}{}{0pt}{}[\\titlerule]",
        "\\titlespacing{\\section}{0pt}{10pt}{4pt}",
        "",
        "\\begin{document}",
        "",
    ]

    if contact.name:
        out.append(f"\\begin{{center}}\n{{\\LARGE\\bfseries {_tex(contact.name)}}}\\\\[2pt]")
        details = [d for d in (contact.email, contact.phone, contact.location) if d]
        if details:
            # Escape each part, then join with the separator — never escape the
            # separator and unescape it afterwards, which silently mangles a
            # real dollar sign somebody typed in a salary line.
            out.append(" $\\cdot$ ".join(_tex(part) for part in details) + "\\\\")
        if contact.links:
            out.append(" $\\cdot$ ".join(_tex(link) for link in contact.links))
        out.append("\\end{center}\n")

    for block in document.sections:
        if block.kind == "header":
            continue
        if block.title:
            out.append(f"\\section*{{{_tex(block.title)}}}")
        for node in block.loose_nodes:
            out.append(_tex(node.text) + "\n")
        for entry in block.entries:
            heading = _entry_heading(entry)
            if heading:
                out.append(f"\\textbf{{{_tex(heading)}}}\\\\")
            if entry.bullets:
                out.append("\\begin{itemize}")
                out += [f"  \\item {_tex(node.text)}" for node in entry.bullets]
                out.append("\\end{itemize}")
        out.append("")

    out.append("\\end{document}")
    return "\n".join(out) + "\n"


# ═══════════════════════════════════════════════════════════════════════════
# Plain text and Markdown
# ═══════════════════════════════════════════════════════════════════════════


def render_text(document: CVDocument) -> str:
    """Plain text, laid out so it survives being pasted into a form field.

    Applicant tracking systems parse this better than anything else, and a
    surprising number of applications are still a textarea.
    """
    contact = document.contact
    out: list[str] = []

    if contact.name:
        out += [contact.name.upper(), ""]
    details = [d for d in (contact.email, contact.phone, contact.location) if d]
    if details:
        out.append(" | ".join(details))
    out += list(contact.links)
    if out:
        out.append("")

    for block in document.sections:
        if block.kind == "header":
            continue
        if block.title:
            out += [block.title.upper(), "-" * len(block.title), ""]
        for node in block.loose_nodes:
            out += [node.text, ""]
        for entry in block.entries:
            heading = _entry_heading(entry)
            if heading:
                out.append(heading)
            out += [f"- {node.text}" for node in entry.bullets]
            if entry.bullets or heading:
                out.append("")

    return "\n".join(out).rstrip() + "\n"


def render_markdown(document: CVDocument) -> str:
    contact = document.contact
    out: list[str] = []

    if contact.name:
        out += [f"# {contact.name}", ""]
    details = [d for d in (contact.email, contact.phone, contact.location) if d]
    if details:
        out += [" · ".join(details), ""]
    if contact.links:
        out += [" · ".join(contact.links), ""]

    for block in document.sections:
        if block.kind == "header":
            continue
        if block.title:
            out += [f"## {block.title}", ""]
        for node in block.loose_nodes:
            out += [node.text, ""]
        for entry in block.entries:
            heading = _entry_heading(entry)
            if heading:
                out += [f"**{heading}**", ""]
            out += [f"- {node.text}" for node in entry.bullets]
            if entry.bullets:
                out.append("")

    return "\n".join(out).rstrip() + "\n"


# ═══════════════════════════════════════════════════════════════════════════


def _heading_text(title: str, transform: str) -> str:
    """Apply the original's heading case, so a rebuild reads like the source.

    A CV that shouted EXPERIENCE should keep shouting it; one that wrote
    "Experience" should not start.
    """
    if transform == "upper":
        return title.upper()
    if transform == "title":
        return title.title()
    return title


def _entry_heading(entry) -> str:
    left = " · ".join(part for part in (entry.role, entry.org) if part)
    if entry.location:
        left = f"{left}, {entry.location}" if left else entry.location
    when = " – ".join(part for part in (entry.start, entry.end) if part)
    if left and when:
        return f"{left} — {when}"
    return left or when


__all__ = ["REBUILD_NOTE", "render", "render_docx", "render_markdown", "render_tex", "render_text"]
